"""
Word Document Parser — Extracts transaction data from .docx files.
Handles tables (police complaint attachments) and paragraph text (narratives).
"""
import re
import logging
from typing import List, Tuple, Dict, Optional
from datetime import datetime

from models.schemas.ingestion_multi import NormalizedTransaction, ParseError

logger = logging.getLogger(__name__)


async def parse_docx(
    file_path: str,
    original_filename: str
) -> Tuple[List[NormalizedTransaction], List[ParseError]]:
    """
    Extracts transactions from Word documents.
    Strategy 1: Parse tables (police complaint attachments often have tables)
    Strategy 2: Parse paragraph text with regex
    """
    transactions: List[NormalizedTransaction] = []
    errors: List[ParseError] = []

    try:
        from docx import Document
    except ImportError:
        errors.append(ParseError(
            filename=original_filename,
            reason="python-docx not installed. Run: pip install python-docx",
            severity="error"
        ))
        return transactions, errors

    try:
        doc = Document(file_path)

        # Strategy 1: Tables
        for table_idx, table in enumerate(doc.tables):
            header_row: Optional[List[str]] = None

            for row_idx, row in enumerate(table.rows):
                cells = [cell.text.strip() for cell in row.cells]

                # Detect header row
                cells_text = ' '.join(cells).lower()
                if row_idx == 0 or any(kw in cells_text for kw in
                                       ['date', 'amount', 'transaction', 'debit', 'credit', 'particular']):
                    header_row = [c.lower().strip() for c in cells]
                    continue

                if header_row and len(cells) >= len(header_row):
                    row_dict = {}
                    for i, cell in enumerate(cells):
                        if i < len(header_row):
                            row_dict[header_row[i]] = cell

                    txn = _extract_from_word_row(row_dict, original_filename, table_idx, row_idx)
                    if txn:
                        transactions.append(txn)

        # Strategy 2: Paragraph text regex (if no tables found meaningful data)
        if not transactions:
            full_text = "\n".join(para.text for para in doc.paragraphs if para.text.strip())

            if full_text.strip():
                transactions = _parse_narrative_text(full_text, original_filename)
                for txn in transactions:
                    txn.confidence = 0.55
                    txn.parser_notes = "Extracted from Word document narrative text"

        if not transactions:
            errors.append(ParseError(
                filename=original_filename,
                reason="No transaction data found in Word document",
                severity="warning"
            ))

    except Exception as e:
        logger.error(f"Word document parse error for {original_filename}: {e}", exc_info=True)
        errors.append(ParseError(
            filename=original_filename,
            reason=f"Word document parsing error: {str(e)}",
            severity="error"
        ))

    return transactions, errors


def _extract_from_word_row(
    row: Dict[str, str],
    source_file: str,
    table_idx: int,
    row_idx: int
) -> Optional[NormalizedTransaction]:
    """Extract a NormalizedTransaction from a Word table row dict."""
    # Find date
    date_val = None
    for key in ['date', 'txn date', 'transaction date', 'dt']:
        if key in row and row[key]:
            date_val = row[key]
            break

    if not date_val:
        return None

    # Find amount
    amount = 0.0
    for key in ['amount', 'debit', 'credit', 'value', 'withdrawal', 'deposit', 'txn amount']:
        if key in row and row[key]:
            cleaned = re.sub(r'[₹,\s]', '', row[key])
            try:
                amt = abs(float(cleaned))
                if amt > 0:
                    amount = amt
                    break
            except ValueError:
                continue

    if amount == 0:
        return None

    # Parse date
    try:
        from services.ingestion.pdf_parser import _parse_indian_date
        timestamp = _parse_indian_date(date_val)
    except Exception:
        timestamp = datetime.utcnow()

    # Find accounts
    from_acc = "UNKNOWN_SENDER"
    to_acc = "UNKNOWN_RECEIVER"
    for key in ['from', 'sender', 'from account', 'debit account', 'source']:
        if key in row and row[key]:
            from_acc = row[key]
            break
    for key in ['to', 'receiver', 'to account', 'credit account', 'beneficiary', 'destination']:
        if key in row and row[key]:
            to_acc = row[key]
            break

    # Find reference
    ref = ""
    for key in ['ref', 'reference', 'utr', 'txn id', 'rrn']:
        if key in row and row[key]:
            ref = row[key]
            break
    if not ref:
        ref = f"DOCX-T{table_idx}R{row_idx}-{int(amount)}"

    # Narration
    narration = ""
    for key in ['narration', 'description', 'particulars', 'remarks', 'details']:
        if key in row and row[key]:
            narration = row[key]
            break

    # Mode
    mode_match = re.search(r'\b(UPI|IMPS|RTGS|NEFT|ATM)\b',
                           f"{narration} {' '.join(row.values())}", re.IGNORECASE)

    return NormalizedTransaction(
        transaction_ref=ref,
        timestamp=timestamp,
        amount=amount,
        from_account=from_acc,
        to_account=to_acc,
        transaction_type=mode_match.group(1).upper() if mode_match else None,
        narration=narration[:300],
        source_file=source_file,
        confidence=0.80,
        parser_notes=f"Word table {table_idx + 1}, row {row_idx + 1}"
    )


def _parse_narrative_text(text: str, source_file: str) -> List[NormalizedTransaction]:
    """Parse amounts and dates from narrative text in complaint letters."""
    transactions = []
    date_pattern = r'(\d{1,2}[/\-]\d{1,2}[/\-]\d{2,4})'
    amount_pattern = r'₹?\s*([\d,]+\.?\d{0,2})'
    account_pattern = r'(?:A/?c|Account|Acc)[:\s#]*(\d{8,18})'

    lines = text.split('\n')
    for i, line in enumerate(lines):
        if len(line.strip()) < 15:
            continue

        date_match = re.search(date_pattern, line)
        amount_matches = re.findall(amount_pattern, line)

        if date_match and amount_matches:
            try:
                from services.ingestion.pdf_parser import _parse_indian_date
                timestamp = _parse_indian_date(date_match.group(1))
                amounts = [abs(float(a.replace(',', ''))) for a in amount_matches]
                amounts = [a for a in amounts if a > 0]
                if not amounts:
                    continue
                amount = max(amounts)

                account_match = re.search(account_pattern, line)
                upi_match = re.search(r'([a-zA-Z0-9._\-]+@[a-zA-Z]+)', line)

                txn = NormalizedTransaction(
                    transaction_ref=f"DOCX-L{i}-{int(amount)}",
                    timestamp=timestamp,
                    amount=amount,
                    from_account=account_match.group(1) if account_match else "UNKNOWN_SENDER",
                    to_account="UNKNOWN_RECEIVER",
                    upi_id=upi_match.group(1) if upi_match else None,
                    narration=line.strip()[:300],
                    source_file=source_file,
                    confidence=0.55,
                    parser_notes=f"Narrative text, line {i + 1}"
                )
                transactions.append(txn)
            except Exception:
                continue

    return transactions
